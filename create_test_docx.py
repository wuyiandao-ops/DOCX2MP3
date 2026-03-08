from docx import Document
import os

def create_test_file():
    doc = Document()
    
    # 1. 標題
    doc.add_heading('DOCX 轉 MP3 功能測試文件', 0)
    
    # 2. 破音字/字典替換測試區
    doc.add_heading('1. 破音字與字典替換測試', level=1)
    doc.add_paragraph('這是一個專門測試「發音字典」替換功能的段落。')
    doc.add_paragraph('測試詞彙：般若波羅蜜多心經。')
    doc.add_paragraph('預期行為：如果您在發音字典中設定「般若」替換為「波惹」，合成後的語音應讀作「波惹」。')
    doc.add_paragraph('其他常見測試詞彙：重慶、長沙、銀行。')
    
    # 3. 長文本智慧切分測試區 (模擬超過 10,000 字)
    doc.add_heading('2. 長文本智慧切分測試 (超過 10,000 字)', level=1)
    doc.add_paragraph('本區段包含重複的文字，總字數將超過 15,000 字，用以測試系統是否會自動切分為兩個段落。')
    
    # 重複文字以達到字數目標 (一段約 50 字，重複 300 次約 15,000 字)
    long_content = "這是一個長文本測試段落。我們正在驗證智慧切分引擎是否能正確識別句點，並在接近一萬字的地點自動將文件切開，以確保語音合成不會因為字數過多而崩潰。"
    for i in range(1, 301):
        doc.add_paragraph(f"第 {i} 組測試文字：{long_content}")
    
    # 4. 結尾
    doc.add_paragraph('--- 測試文件結束 ---')
    
    # 儲存檔案
    file_name = 'test_sample.docx'
    doc.save(file_name)
    print(f"✅ 測試文件 '{file_name}' 已成功產生！")
    print(f"檔案路徑: {os.path.abspath(file_name)}")

if __name__ == "__main__":
    create_test_file()
